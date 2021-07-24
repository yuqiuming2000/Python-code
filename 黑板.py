from docx import Document
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    # p._p = p._element = None
    paragraph._p = paragraph._element = None



doc= Document('昌碧.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[30].text)
print(doc.paragraphs[13].text)
for i in reversed(range(13,31)):
    m=doc.paragraphs[i]
    print(m.text)
    delete_paragraph(m)

    # if ''.join(p.text.split(' ')).lower()=='header_keyword':
    #     break
# for p in doc.paragraphs:  
#     if p.text.lower()=='': # 删除word中在开始部分的空白段落
#         delete_paragraph(p)
#     else:
#         break


doc.save('昌碧删除后.docx')