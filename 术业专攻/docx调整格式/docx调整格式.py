import os
from docx import Document
from docxcompose.composer import Composer
from docx import Document as Document_compose
from docx.shared import Inches,Mm, Pt,Length,RGBColor,Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    # p._p = p._element = None
    paragraph._p = paragraph._element = None
path='术业专攻\docx调整格式\调整前/'
new_path='术业专攻\docx调整格式\调整后/'
files=os.listdir(path)
for file in files:
    file_path=path+file
    new_file_path=new_path+'格式调整后'+file
    doc = Document(file_path)
    print(file)
    print(len(doc.paragraphs))
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.bold = False
            # run.font.italic = True
            # run.font.underline = True
            # run.font.strike = True
            # run.font.shadow = True
            run.font.size = Pt(15)
            # run.font.color.rgb = RGBColor(255,0,255)
            run.font.name = "仿宋"
            r = run._element.rPr.rFonts
            r.set(qn("w:eastAsia"),'仿宋')
        
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2) 
    for i in reversed(range(13,31)):
        m=doc.paragraphs[i]
        print(m.text)
        delete_paragraph(m)

    print(len(doc.paragraphs))

    for run in doc.paragraphs[1].runs:
        run.font.bold = True
        run.font.size = Pt(22)
        run.font.name = "宋体"
        r = run._element.rPr.rFonts
        r.set(qn("w:eastAsia"),'宋体')
    doc.paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.CENTER  
    doc.paragraphs[11].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.paragraphs[12].alignment = WD_ALIGN_PARAGRAPH.RIGHT 
    sections = doc.sections
    for section in sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
    
    doc.save(new_file_path)